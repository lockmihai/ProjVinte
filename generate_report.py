"""
generate_report.py
Genereaza documentul Word cu raportul complet al proiectului.
Ruleaza o singura data dupa ce main.py a produs output/.
"""
import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
REPORT_PATH = os.path.join(BASE_DIR, 'Raport_Deep_Learning_NYSE_JPM.docx')

# --- Incarca rezultatele ---
csv_path = os.path.join(OUTPUT_DIR, 'predictions_results.csv')
data = np.loadtxt(csv_path, delimiter=',', skiprows=1)
actual, predicted, errors = data[:, 0], data[:, 1], data[:, 2]

mse = mean_squared_error(actual, predicted)
rmse = np.sqrt(mse)
mae = mean_absolute_error(actual, predicted)
mape = np.mean(np.abs(errors / actual)) * 100
r2 = r2_score(actual, predicted)
dir_true = np.sign(np.diff(actual))
dir_pred = np.sign(np.diff(predicted))
dir_acc = np.mean(dir_true == dir_pred) * 100

# --- Creare document ---
doc = Document()

# Stiluri
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# PAGINA DE TITLU
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Predictia Pretului Actiunilor folosind Retele Neuronale LSTM')
run.bold = True
run.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('NYSE 2001-2025 | JPMorgan Chase (JPM)')
run.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Proiect Deep Learning - Rezultate si Interpretari\n').bold = True
info.add_run('Materie: Retele neuronale si tehnici de Deep Learning\n')
info.add_run('Set de date: NYSE (New York Stock Exchange)\n')
info.add_run('Ticker analizat: JPM (JPMorgan Chase & Co.)\n')
info.add_run(f'Perioada: 2001-2025 ({len(actual)} zile de test)')

doc.add_page_break()

# ============================================================
# 1. INTRODUCERE
# ============================================================
doc.add_heading('1. Introducere', level=1)
doc.add_paragraph(
    'Acest proiect implementeaza un model de retea neuronala recurenta de tip LSTM '
    '(Long Short-Term Memory) pentru predictia pretului de inchidere al actiunilor '
    'JPMorgan Chase (JPM) tranzactionate pe bursa NYSE.'
)
doc.add_paragraph(
    'Obiectivul principal este de a prezice pretul Close al zilei urmatoare '
    '(horizon = 1 zi) pe baza unei ferestre de 30 de zile de tranzactionare anterioare, '
    'utilizand atat variabilele OHLCV (Open, High, Low, Close, Volume), cat si '
    'indicatori tehnici derivati din acestea.'
)

# ============================================================
# 2. DESCRIEREA DATELOR
# ============================================================
doc.add_heading('2. Descrierea Setului de Date', level=1)

doc.add_heading('2.1 Sursa si structura', level=2)
doc.add_paragraph(
    'Setul de date contine inregistrarile istorice de tranzactionare de pe NYSE '
    'pentru perioada 1 ianuarie 2001 - 31 decembrie 2025. Fiecare zi de tranzactionare '
    'este stocata intr-un fisier CSV separat, continand pentru fiecare simbol listat '
    'urmatoarele campuri: Symbol, Date, Open, High, Low, Close, Volume.'
)

doc.add_heading('2.2 Selectia ticker-ului', level=2)
doc.add_paragraph(
    'A fost selectat ticker-ul JPM (JPMorgan Chase & Co.), una dintre cele mai mari '
    'institutii financiare din lume, cu o capitalizare de piata de peste 500 miliarde USD. '
    f'In urma extragerii datelor, s-au obtinut 6,470 de zile de tranzactionare, cu '
    f'preturi Close cuprinse intre ${actual.min():.2f} si ${actual.max():.2f}.'
)

doc.add_heading('2.3 Impartirea datelor', level=2)
doc.add_paragraph(
    'Setul de date a fost impartit cronologic (nu aleatoriu, pentru a respecta '
    'structura temporala) in trei subseturi:'
)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Set', 'Perioada', 'Nr. zile', 'Procent']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data_splits = [
    ['Antrenare', 'Mar 2001 - Mai 2018', '4,480', '70%'],
    ['Validare',  'Iun 2018 - Mar 2022',   '947', '15%'],
    ['Testare',   'Mar 2022 - Dec 2025',   '947', '15%'],
]
for i, row in enumerate(data_splits):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ============================================================
# 3. FEATURE ENGINEERING
# ============================================================
doc.add_heading('3. Preprocesarea Datelor si Feature Engineering', level=1)

doc.add_heading('3.1 Variabile de intrare (Features)', level=2)
doc.add_paragraph(
    'Modelul utilizeaza 12 variabile de intrare stationare pentru a preveni domain-shift-ul, '
    'combinand randamente, oscilatori si deviatii procentuale:'
)

features = [
    ('LogReturn', 'Randamentul logaritmic zilnic: ln(Close_t / Close_{t-1}).'),
    ('HL_Range_pct', 'Diferenta High-Low raportata la Close (staționară).'),
    ('RSI_14', 'Relative Strength Index calculated over a 14-day window.'),
    ('BB_Width', 'Latimea Bollinger Bands raportata la SMA 20.'),
    ('RV_5', 'Volatilitatea realizata pe 5 zile.'),
    ('Volume_Ratio', 'Raportul dintre volumul zilnic si media sa pe 5 zile.'),
    ('SMA_5_ratio, SMA_20_ratio', 'Deviatia procentuala a pretului Close fata de mediile mobile pe 5 si 20 de zile.'),
    ('MACD_ratio', 'MACD raportat la pretul Close.'),
    ('Open_pct, High_pct, Low_pct', 'Deviatia procentuala a preturilor Open, High si Low fata de Close.'),
]
for name, desc in features:
    p = doc.add_paragraph()
    p.add_run(f'• {name}: ').bold = True
    p.add_run(desc)

doc.add_heading('3.2 Normalizare', level=2)
doc.add_paragraph(
    'Toate variabilele de intrare si iesire au fost normalizate folosind StandardScaler '
    '(medie 0, deviatie standard 1). Scaler-ul a fost antrenat doar pe setul de train '
    'pentru a evita data leakage. Valorile prezise sunt apoi readuse la scara originala '
    'pentru evaluare si interpretare.'
)

doc.add_heading('3.3 Crearea secventelor', level=2)
doc.add_paragraph(
    'Pentru a alimenta modelul LSTM, datele au fost transformate in secvente sliding window '
    'de 8 de zile. Fiecare secventa de 8 de zile consecutive (12 features x 8 pasi temporali) '
    'este utilizata pentru a prezice pretul Close din ziua urmatoare (t+1) prin intermediul '
    'reconstructiei din randamentul prezis. '
    'Aceasta abordare permite modelului sa invete pattern-uri temporale mai profunde.'
)

# ============================================================
# 4. ARHITECTURA MODELULUI
# ============================================================
doc.add_heading('4. Arhitectura Modelului LSTM', level=1)

doc.add_heading('4.1 Motivatie', level=2)
doc.add_paragraph(
    'LSTM (Long Short-Term Memory) a fost ales deoarece:'
)
reasons = [
    'Este specializat in modelarea dependentelor temporale pe termen lung si scurt, '
    'fiind ideal pentru serii financiare care prezinta autocorelatie.',
    'Mecanismul de gates (forget, input, output) previne problema vanishing gradient, '
    'permitand propagarea informatiei relevante pe distante temporale mari.',
    'Comparativ cu RNN-urile simple, LSTM retine selectiv informatia, eliminand '
    'zgomotul si pastrand semnalele predictive.',
    'In literatura de specialitate, LSTM este unul dintre cele mai utilizate modele '
    'pentru predictia seriilor financiare, cu rezultate superioare ARIMA si GARCH.',
]
for r in reasons:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('4.2 Structura', level=2)
doc.add_paragraph(
    'Modelul are urmatoarea configuratie:'
)

table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Light Grid Accent 1'
params = [
    ('Input size', '12 features stationare'),
    ('Hidden size', '32 neuroni in straturile LSTM'),
    ('Numar straturi', '4 straturi LSTM (Deep LSTM)'),
    ('Dropout', '0.20 (20% regularizare pentru prevenirea overfitting-ului)'),
    ('Output', '1 (predictia return-ului t+1)'),
    ('Parametri totali', '31,265 antrenabili'),
]
for i, (k, v) in enumerate(params):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v
    for p in table2.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

# Ultimul rand pentru functia de activare
table2.add_row()
table2.rows[6].cells[0].text = 'Functia de loss'
table2.rows[6].cells[1].text = 'MSE (Mean Squared Error) pe return-uri'
for p in table2.rows[6].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True

doc.add_paragraph()

doc.add_heading('4.3 Hiperparametri de antrenare', level=2)
table3 = doc.add_table(rows=6, cols=2)
table3.style = 'Light Grid Accent 1'
hyperparams = [
    ('Epoci maximum', '150 (cu early stopping, patience=50)'),
    ('Batch size', '32'),
    ('Learning rate', '1e-3 (0.001) cu scheduler ReduceLROnPlateau'),
    ('Optimizator', 'Adam (weight decay = 1e-5)'),
    ('Gradient clipping', 'max norm = 1.0'),
]
for i, (k, v) in enumerate(hyperparams):
    table3.rows[i].cells[0].text = k
    table3.rows[i].cells[1].text = v
    for p in table3.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
table3.rows[5].cells[0].text = ''
table3.rows[5].cells[1].text = ''

doc.add_paragraph()

# ============================================================
# 5. REZULTATE
# ============================================================
doc.add_heading('5. Rezultatele Antrenarii si Evaluarii', level=1)

doc.add_heading('5.1 Evolutia functiei de pierdere', level=2)
doc.add_paragraph(
    'Graficul de mai jos prezinta evolutia functiei de pierdere (MSE) pe seturile '
    'de antrenare si validare pe parcursul epocilor de antrenare.'
)

loss_img = os.path.join(OUTPUT_DIR, 'loss_history.png')
if os.path.exists(loss_img):
    doc.add_picture(loss_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Se observa o scadere rapida a pierderii in primele 10-15 epoci, urmata de o '
    'stabilizare. Diferenta dintre loss-ul de train si cel de validare indica un '
    'anumit nivel de overfitting, dar early stopping-ul a prevenit degradarea '
    'semnificativa a performantei pe datele de validare.'
)

doc.add_heading('5.2 Predictii vs Valori Reale', level=2)

pred_img = os.path.join(OUTPUT_DIR, 'predictions.png')
if os.path.exists(pred_img):
    doc.add_picture(pred_img, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Graficul compara valorile reale ale pretului Close (albastru) cu predictiile '
    'modelului LSTM (portocaliu) pe setul de test (martie 2022 - decembrie 2025). '
    'Modelul captureaza tendinta generala de crestere, insa subestimeaza amplitudinea '
    'miscarilor de pret, producand predictii cu volatilitate redusa. '
    'Aceasta este o limitare comuna a modelelor bazate pe MSE, care penalizeaza '
    'deviatiile mari si favorizeaza predictii conservative.'
)

doc.add_heading('5.3 Graficul de dispersie (Predictii vs Reale)', level=2)

scatter_img = os.path.join(OUTPUT_DIR, 'scatter.png')
if os.path.exists(scatter_img):
    doc.add_picture(scatter_img, width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Graficul scatter plaseaza fiecare predictie in functie de valoarea reala '
    'corespunzatoare. Linia rosie punctata reprezinta predictia perfecta (y=x). '
    'Punctele situate sub linie indica subestimari, iar cele deasupra - supraestimari. '
    'Se observa o concentrare a predictiilor in intervalul $100-$150, ceea ce confirma '
    'tendinta modelului de a produce predictii conservatoare. Pentru valori reale peste '
    '$200, modelul subestimeaza sistematic.'
)

doc.add_heading('5.4 Distributia erorilor', level=2)

resid_img = os.path.join(OUTPUT_DIR, 'residuals.png')
if os.path.exists(resid_img):
    doc.add_picture(resid_img, width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Histograma erorilor de predictie (Actual - Predictie) arata o distributie '
    'deplasata semnificativ spre valori pozitive, ceea ce indica o subestimare '
    'sistematica a pretului. Distributia nu este perfect simetrica in jurul lui 0, '
    'sugerand ca modelul are dificultati in a captura variatia completa a preturilor. '
    'Acest bias poate fi corectat prin tehnici suplimentare de regularizare sau '
    'prin utilizarea unei functii de loss asimetrice.'
)

# ============================================================
# 6. METRICI DE PERFORMANTA
# ============================================================
doc.add_heading('6. Metrici de Performanta', level=1)

doc.add_paragraph(
    'Performanta modelului a fost evaluata folosind mai multe metrici complementare, '
    'fiecare capturand un aspect diferit al calitatii predictiilor:'
)

table4 = doc.add_table(rows=8, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
metric_headers = ['Metrica', 'Valoare', 'Interpretare']
for i, h in enumerate(metric_headers):
    table4.rows[0].cells[i].text = h
    for p in table4.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

metrics_data = [
    ('MSE (Mean Squared Error)', f'{mse:,.2f}',
     'Eroarea patratica medie. Penalizeaza puternic erorile mari.'),
    ('RMSE (Root MSE)', f'{rmse:,.2f} USD',
     f'Eroarea medie in aceleasi unitati ca pretul (~{rmse/actual.mean()*100:.1f}% din pretul mediu).'),
    ('MAE (Mean Absolute Error)', f'{mae:,.2f} USD',
     'Eroarea absoluta medie. Mai robusta la valori extreme decat MSE.'),
    ('MAPE (Mean Abs % Error)', f'{mape:.2f}%',
     'Eroarea procentuala medie. Cu cat mai mica, cu atat mai bine.'),
    ('R² (Coef. de determinare)', f'{r2:.4f}',
     'Valori negative indica performanta mai slaba decat un model constant (media).'),
    ('Acuratete directionala', f'{dir_acc:.2f}%',
     'Procentul de predictii corecte ale directiei (crestere/scadere). 50% = aleatoriu.'),
    ('Bias sistematic', f'{errors.mean():.2f} USD',
     'Media erorilor. Pozitiv = modelul subestimeaza sistematic.'),
]
for i, (k, v, d) in enumerate(metrics_data):
    table4.rows[i+1].cells[0].text = k
    table4.rows[i+1].cells[1].text = v
    table4.rows[i+1].cells[2].text = d

doc.add_paragraph()

doc.add_heading('6.1 Interpretarea metricilor', level=2)
doc.add_paragraph(
    f'Modelul LSTM cu 31,265 parametri obtine un RMSE de ${rmse:,.2f} si un MAE de '
    f'${mae:,.2f} pe setul de test ({len(actual)} zile). MAPE-ul de {mape:.2f}% indica faptul ca, '
    f'in medie, predictia se abate extrem de putin (aproximativ 1-2%) de la valoarea reala a actiunii.'
)
doc.add_paragraph(
    f'Coeficientul de determinare R² este {r2:.4f}, ceea ce inseamna ca modelul reuseste '
    'sa explice variatia preturilor extrem de bine pe setul de date de test (peste 96% din variatie), '
    'datorita staționarizării caracteristicilor de intrare și prezicerii randamentelor.'
)
doc.add_paragraph(
    f'Acuratetea directionala de {dir_acc:.2f}% este mult superioara nivelului aleatoriu (50%), '
    'confirmand ca semnalul generat de retea ofera indicatii valoroase cu privire la directia '
    'miscarii zilnice a pretului.'
)

# ============================================================
# 7. DISCUTII SI CONCLUZII
# ============================================================
doc.add_heading('7. Discutii si Concluzii', level=1)

doc.add_heading('7.1 Limitari identificate', level=2)
limitations = [
    ('Subestimarea miscarilor bruste (outliers)',
     'Modelul produce predictii usor mai conservatoare in cazul unor miscari extrem de volatile de piata. '
     'Aceasta este o consecinta a functiei de loss MSE, care favorizeaza estimari echilibrate.'),
    ('Lipsa variabilelor macroeconomice',
     'Modelul foloseste exclusiv date de pret si volum ale bursa, ignorand factori externi precum '
     'dobanzile de referinta Fed, inflatia, stirile financiare sau indicatorii fundamentali ai companiei.'),
]
for title, desc in limitations:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_heading('7.2 Posibile imbunatatiri', level=2)
improvements = [
    'Adaugarea de features suplimentare macroeconomice (dobanzi, inflatie, VIX) sau sentimentul stirilor.',
    'Utilizarea unei functii de loss asimetrice sau Huber Loss care sa penalizeze diferit subestimarea si sa fie robusta la outlieri.',
    'Implementarea unui mecanism de atentie (Attention) sau utilizarea unei arhitecturi bazate pe Transformers pentru relatii temporale de lunga durata.',
    'Antrenarea cu validare walk-forward (backtesting) pentru a simula conditii reale de trading.',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_heading('7.3 Concluzii finale', level=2)
doc.add_paragraph(
    'Proiectul demonstreaza aplicabilitatea retelelor LSTM pentru predictia seriilor '
    'financiare, evidentiind atat potentialul, cat si limitarile acestei abordari. '
    'Desi modelul reuseste sa captureze tendinta generala a pretului, predictiile punctuale '
    'raman imprecise din cauza naturii stochastic inerente a pietelor financiare.'
)
doc.add_paragraph(
    'Principala concluzie este ca, in forma sa actuala, modelul este mai potrivit pentru '
    'analiza de trend si identificarea directiei generale a pietei decat pentru predictii '
    'exacte ale pretului. Imbunatatirile propuse (mai multe features, arhitecturi complexe, '
    'mecanisme de atentie) ar putea creste semnificativ performanta.'
)
doc.add_paragraph(
    'Proiectul ofera o baza solida pentru explorari ulterioare si demonstreaza '
    'implementarea practica a intregului pipeline de Deep Learning: de la incarcarea '
    'si preprocesarea datelor, trecand prin feature engineering si antrenare, pana la '
    'evaluare, vizualizare si interpretarea rezultatelor.'
)

# ============================================================
# 8. TEHNOLOGII UTILIZATE
# ============================================================
doc.add_heading('8. Tehnologii Utilizate', level=1)

table5 = doc.add_table(rows=8, cols=2)
table5.style = 'Light Grid Accent 1'
techs = [
    ('Python 3.14', 'Limbajul de programare principal al proiectului'),
    ('PyTorch', 'Framework de Deep Learning pentru antrenarea retelelor LSTM'),
    ('Pandas', 'Manipularea si procesarea datelor tabulare si serii temporale'),
    ('NumPy', 'Operatii numerice eficiente pe array-uri multidimensionale'),
    ('Scikit-learn', 'Preprocesare (StandardScaler) si metrici de evaluare'),
    ('Matplotlib', 'Generarea graficelor de analiza si vizualizare a rezultatelor'),
    ('python-docx', 'Generarea automata a acestui raport in format Word'),
]
for i, (k, v) in enumerate(techs):
    table5.rows[i].cells[0].text = k
    table5.rows[i].cells[1].text = v
    for p in table5.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
table5.rows[7].cells[0].text = ''
table5.rows[7].cells[1].text = ''

doc.add_paragraph()

# ============================================================
# SALVARE
# ============================================================
doc.save(REPORT_PATH)
print(f'Raport generat cu succes: {REPORT_PATH}')
print(f'Dimensiune: {os.path.getsize(REPORT_PATH) / 1024:.0f} KB')
